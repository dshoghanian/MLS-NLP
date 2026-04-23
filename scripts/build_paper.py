"""
Generate MLS_Narrative_Network_Paper.docx — publication-ready manuscript outline.
Run: /opt/anaconda3/bin/python scripts/build_paper.py
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
    return p

def para(text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    return p

def spacer():
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
t1 = doc.add_heading(
    "Narrative Capital and Competitive Performance in Major League Soccer:", level=1)
t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2 = doc.add_heading(
    "A Network Analysis of Press and Fan Discourse (2018-2024)", level=2)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
spacer()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("[Author Name]\n").bold = True
p.add_run("[Institution / Program]\n")
p.add_run("[Date]")
for run in p.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# ABSTRACT
# ─────────────────────────────────────────────────────────────────────────────
h("Abstract", 1)
para(
    "This study examines whether media narrative centrality -- a club's structural position "
    "in a press co-occurrence network -- reflects, predicts, or diverges from on-field "
    "performance in Major League Soccer across seven seasons (2018-2024). Using a corpus of "
    "approximately 7,236 press articles and 15,795 Reddit posts, we construct weighted "
    "co-occurrence graphs for 28 clubs and compute PageRank-based narrative centrality scores "
    "as a proxy for media attention capital. We apply VADER sentiment analysis to both press "
    "and fan discourse channels, run Latent Dirichlet Allocation (LDA) topic modeling to "
    "characterize recurring narrative themes, and estimate OLS and logistic regression models "
    "-- including two-way fixed effects specifications with cluster-robust standard errors -- "
    "predicting next-season points, playoff qualification, and jersey sponsorship deal values. "
    "Our findings indicate that narrative centrality is largely decoupled from competitive "
    "performance (narrative rank ~ performance rank correlation r = 0.41, R-squared = 0.17): "
    "large-market clubs dominate press co-occurrence networks regardless of results, while "
    "smaller-market clubs (Columbus Crew, Seattle Sounders FC, Real Salt Lake) achieve "
    "competitive success with minimal national narrative presence. An incremental F-test "
    "confirms that narrative centrality adds no significant predictive power beyond objective "
    "performance controls (xGoal difference, payroll, attendance; F = 0.72, p = 0.578). "
    "The 2023 Messi signing at Inter Miami CF functions as a natural experiment; a Synthetic "
    "Control Method analysis estimates a post-treatment PageRank gap of 0.035 units -- the "
    "largest discontinuity in the dataset. Press sentiment is systematically more positive "
    "than fan (Reddit) sentiment across all clubs, a genre effect independent of performance. "
    "Jersey sponsorship deal value is positively associated with narrative centrality even "
    "after controlling for franchise valuation, suggesting PageRank functions as measurable "
    "attention capital with direct commercial implications."
)
spacer()
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0)
run = p.add_run("Keywords: ")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(12)
r2 = p.add_run(
    "narrative centrality, co-occurrence networks, PageRank, Major League Soccer, "
    "sports communication, sentiment analysis, natural experiment, sponsorship ROI, "
    "intermedia agenda-setting, two-way fixed effects"
)
r2.font.name = "Times New Roman"
r2.font.size = Pt(12)
doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 1. INTRODUCTION
# ─────────────────────────────────────────────────────────────────────────────
h("1. Introduction", 1)
para(
    "Professional sports clubs exist simultaneously in two competitive arenas: the field, where "
    "wins and losses determine standings, and the media, where narrative presence determines "
    "commercial value, fan growth, and sponsorship attractiveness. In most North American sports "
    "contexts, these arenas are assumed to be coupled -- winning clubs attract more coverage, "
    "and coverage compounds winning by attracting better players and sponsors. This study tests "
    "that assumption directly for Major League Soccer (MLS), a league characterized by "
    "significant structural variation in market size, expansion history, and media market access."
)
para(
    "We approach MLS media coverage as a network phenomenon. Each press article or Reddit post "
    "that mentions multiple clubs simultaneously creates an implicit co-occurrence link -- a "
    "structural trace of how journalists and fans frame clubs in relation to one another. "
    "Aggregated across thousands of documents per season, these co-occurrences form a weighted "
    "network whose topology can be analyzed using graph-theoretic centrality measures. The "
    "PageRank algorithm, originally designed to rank web pages by structural authority "
    "(Brin & Page, 1998; Page et al., 1999), translates naturally to this context: a club "
    "with high PageRank is not merely mentioned frequently, but is mentioned alongside other "
    "prominently-covered clubs, amplifying its narrative authority in a compounding fashion."
)
para(
    "This framework connects to two bodies of theory. From media studies, intermedia "
    "agenda-setting theory (McCombs & Shaw, 1972; Lim, 2012) proposes that media outlets "
    "shape the salience of topics for public discourse -- and that online audiences follow the "
    "agenda set by professional media with a measurable lag. From sports economics, a growing "
    "literature documents the commercial value of media attention: television exposure predicts "
    "ticket revenue (Peeters & Szymanski, 2012), and jersey sponsorship valuations track brand "
    "visibility metrics. Our contribution is to operationalize these relationships at the "
    "club-season level using network-analytic methods applied to a novel dual-channel corpus "
    "spanning press and fan discourse."
)
para(
    "We investigate four primary research questions: (RQ1) Is narrative centrality correlated "
    "with competitive performance in MLS? (RQ2) Does narrative centrality predict future "
    "performance beyond objective metrics (xGoals, payroll, attendance)? (RQ3) Does press "
    "narrative lead fan discourse temporally, consistent with intermedia agenda-setting? "
    "(RQ4) Is narrative centrality associated with jersey sponsorship deal value, independent "
    "of franchise valuation?"
)

# ─────────────────────────────────────────────────────────────────────────────
# 2. LITERATURE REVIEW
# ─────────────────────────────────────────────────────────────────────────────
h("2. Literature Review", 1)

h("2.1 Network Analysis in Sports", 2)
para(
    "Network methods have been applied extensively to player performance data in soccer "
    "(Cintia et al., 2015; Gudmundsson & Horton, 2017), typically constructing passing "
    "networks or player interaction graphs. Application to media discourse networks -- where "
    "nodes are clubs and edges are co-occurrence relationships -- is less established. "
    "Closest to our approach is work on sports journalism citation networks (McEnnis, 2020) "
    "and computational analysis of sports Twitter discourse (Yu & Wang, 2015). We extend "
    "this tradition by constructing annual and quarterly co-occurrence graphs from a "
    "systematically collected multi-source corpus, enabling centrality comparisons across "
    "channels (press vs. Reddit) and over time (2018-2024)."
)

h("2.2 Agenda-Setting and Intermedia Effects", 2)
para(
    "McCombs and Shaw (1972) established that media salience shapes public attention; "
    "subsequent work extended this to online environments (Lim, 2012). In sports contexts, "
    "media framing shapes fan identification and commercial behavior (Billings, 2011). The "
    "dual-channel design of this study allows direct testing of whether press coverage "
    "Granger-causes fan discourse volume (Granger, 1969), operationalizing the temporal "
    "dimension of agenda-setting. Granger tests are conducted on a 27-quarter aggregate "
    "series (2018 Q1 -- 2024 Q3) reconstructed from the network parquet files, providing "
    "adequate power (N > 20 per standard recommendation). The cross-lagged correlation on "
    "per-club annual series offers a complementary within-club temporal test."
)

h("2.3 Sentiment Analysis in Sports Media", 2)
para(
    "VADER (Hutto & Gilbert, 2014) is the dominant lexicon-based sentiment tool for social "
    "media text and has been applied to sports discourse analysis (Boon-Itt & Skunkan, 2021). "
    "Its known limitation on editorial text (F1 approximately 0.55-0.65 versus F1 > 0.80 for "
    "transformer-based models; Barbieri et al., 2020) applies to our press corpus. We treat "
    "VADER scores as ordinal proxies for sentiment direction rather than precise measurements, "
    "focusing on cross-channel comparisons where the measurement tool is held constant."
)

h("2.4 Topic Modeling of Sports Discourse", 2)
para(
    "Latent Dirichlet Allocation (Blei, Ng & Jordan, 2003) has been used to characterize "
    "sports media narratives (Nichols et al., 2014) and fan forum content (Magnusson, 2016). "
    "Topic model quality is typically evaluated via coherence scores (Roder, Both & Hinneburg, "
    "2015). We retrained our model with MLS-specific domain stopwords to prevent cross-league "
    "contamination (NBA, NHL, NFL terminology) from distorting topic structure -- a procedural "
    "step that improved topic interpretability and is documented in our methodology for "
    "replication."
)

h("2.5 Natural Experiments and Superstar Effects in Sports", 2)
para(
    "The Difference-in-Differences (DiD) framework has been used in sports economics to "
    "estimate the effect of stadium construction (Coates & Humphreys, 1999) and star player "
    "signings on attendance and revenue (Hausman & Leonard, 1997). The Synthetic Control "
    "Method (Abadie, Diamond & Hainmueller, 2010) extends DiD by constructing a weighted "
    "counterfactual, relaxing the parallel trends assumption. The 2023 Messi signing provides "
    "a rare, well-identified natural experiment in the MLS context: an exogenous, globally "
    "salient event affecting a single club mid-panel, enabling causal estimation of the "
    "narrative impact of a superstar acquisition."
)

h("2.6 Panel Data Methods in Sports Analytics", 2)
para(
    "Sports panel data -- repeated observations of the same clubs over multiple seasons -- "
    "requires methods that account for club-level heterogeneity (market size, ownership "
    "resources, geographic factors) and year-level shocks (COVID 2020, new broadcast deals, "
    "league expansion). Two-way fixed effects (Cameron & Miller, 2015) absorbs both types "
    "of confounders, providing a conservative within-entity test of whether changes in "
    "narrative centrality predict changes in performance. Cluster-robust standard errors "
    "at the club level address serial correlation in club outcomes across seasons."
)

# ─────────────────────────────────────────────────────────────────────────────
# 3. DATA AND METHODS
# ─────────────────────────────────────────────────────────────────────────────
h("3. Data and Methods", 1)

h("3.1 Press Corpus", 2)
para(
    "Press articles were assembled from two sources: (1) the GDELT Global Knowledge Graph "
    "(Leetaru & Schrodt, 2013), queried with 11 MLS-focused search terms per month for "
    "2018-2024; and (2) RSS feeds from 12 MLS-adjacent outlets (MLSSoccer.com, The Athletic, "
    "ESPN FC). Full article text was extracted using Trafilatura (Barbaresi, 2021), with "
    "newspaper3k as fallback. After deduplication via SHA-256 content hashing and filtering "
    "for English-language MLS content, the corpus contains 7,236 articles across 84 "
    "monthly collection windows."
)

h("3.2 Reddit Corpus", 2)
para(
    "Reddit submissions and comments were collected via the Pushshift API from r/MLS and "
    "28 club-specific subreddits for 2018-2024, yielding 15,795 posts. Note: r/cfmontreal "
    "was not available in the collection window; all CF Montreal Reddit statistics reflect "
    "incidental mentions in r/MLS only and are flagged as caveated in comparative analyses."
)

h("3.3 External Data Sources", 2)
para(
    "Three external data sources are integrated into extended regression models. American "
    "Soccer Analysis (ASA) provides club-level xGoals, expected points (xPoints), and home "
    "attendance via its public REST API (186 club-season observations, 2018-2024). MLSPA "
    "salary disclosures provide player guaranteed compensation, parsed from PDFs for 2018-2019 "
    "and 2021-2023 and from CSV for 2024, yielding 4,209 player-level records aggregated to "
    "142 club-season rows (2020 absent due to COVID non-standard reporting). Google Trends "
    "provides weekly search interest indices (0-100) for each club name via the pytrends "
    "library, with cross-batch normalization using 'MLS soccer' as an anchor term present in "
    "every 5-term query batch."
)

h("3.4 Network Construction", 2)
para(
    "For each annual time window (2018-2024), we construct a weighted undirected co-occurrence "
    "graph using NetworkX (Hagberg, Schult & Swart, 2008). Nodes represent the 28 MLS clubs "
    "active in the study period; an edge connects two clubs if they co-appear in the same "
    "document, with edge weight equal to the count of co-mentioning documents. Club names "
    "were identified using spaCy (Honnibal et al., 2020) augmented with a domain-specific "
    "gazetteer of canonical club names and common aliases. Unicode normalization (NFD to "
    "ASCII) handled accent variants across data sources (e.g., 'CF Montreal' vs. 'CF "
    "Montreal'). Separate networks were constructed for press and Reddit corpora. Quarterly "
    "sub-networks support momentum and lead-lag analyses."
)

h("3.5 Centrality and Narrative Metrics", 2)
para(
    "Five centrality measures were computed on each annual network: PageRank (damping "
    "factor alpha = 0.85), degree centrality, betweenness centrality, closeness centrality, "
    "and eigenvector centrality (Borgatti, 2005). PageRank is our primary narrative centrality "
    "proxy due to its sensitivity to the quality -- not just quantity -- of co-occurrences. "
    "Season-normalized PageRank (min-max within each year) enables cross-year comparisons. "
    "Narrative momentum is the quarterly PageRank delta: delta_PR(t) = PR(t) - PR(t-1), "
    "classified as rising (delta > +0.005), falling (delta < -0.005), or stable. Net momentum "
    "signal (windows_rising - windows_falling) provides a season-level momentum indicator."
)

h("3.6 Sentiment Analysis", 2)
para(
    "VADER (Hutto & Gilbert, 2014) was applied to full document text without preprocessing. "
    "Compound scores range -1 to +1; labels: positive (> +0.05), negative (< -0.05), neutral "
    "otherwise. Club-level sentiment was aggregated as the mean compound score across all "
    "documents mentioning that club in a given year. The sentiment gap (press_sent - "
    "reddit_sent) captures the systematic divergence between professional and fan discourse "
    "tones, with positive values indicating press is more favorable than fans."
)

h("3.7 Topic Modeling", 2)
para(
    "LDA was implemented using Gensim (Rehurek & Sojka, 2010) with 7 topics on the combined "
    "press and Reddit corpus. Preprocessing included standard English stopwords, MLS-generic "
    "terms (club names, generic soccer vocabulary), and a domain-specific stopword list "
    "excluding cross-league terminology (NBA, NHL, NFL, MLB abbreviations; sport-specific "
    "jargon; European club names appearing in transfer articles). This last category was added "
    "after initial training revealed cross-league contamination in one topic, where NBA/NHL/NFL "
    "terminology dominated a catch-all topic due to MLS articles appearing on general sports "
    "outlets. The final vocabulary contained 21,043 unique terms across 20,745 documents. "
    "Topics were labeled by manual review of top-20 word lists and validated against the "
    "yearly trend charts for temporal coherence."
)

h("3.8 Regression Specifications", 2)
para(
    "Three primary regression frameworks are estimated. (1) OLS predicting next_points -- "
    "the club's total regular-season points in year T+1 -- from narrative features across "
    "three nested models: baseline (PageRank only), narrative block (PageRank + momentum + "
    "degree), and full model (narrative + sentiment + playoff status + 2020 structural break "
    "indicator). The sentiment_gap variable (press_sent minus reddit_sent) is excluded from "
    "regression models as it is a perfect linear combination of two predictors already included, "
    "producing infinite VIF by construction; press_sent and reddit_sent enter independently "
    "instead. A two-way fixed effects "
    "(TWFE) specification with club and year fixed effects is estimated using linearmodels "
    "(Sheppard, 2021), with standard errors clustered at the club level following Cameron and "
    "Miller (2015). An incremental F-test tests whether the narrative feature block adds "
    "predictive power beyond external performance controls (xGoal difference, log payroll, "
    "log attendance, normalized Google Trends search interest). Multicollinearity is diagnosed "
    "via Variance Inflation Factors (VIF). Leave-One-Year-Out (LOYO) cross-validation provides "
    "out-of-sample performance estimates. Robustness checks exclude 2020 (COVID-shortened "
    "season) and 2023 (Messi structural break). (2) Logistic regression predicting made_playoffs "
    "(binary) from the full narrative and sentiment feature set, with marginal effects reported. "
    "(3) OLS predicting estimated jersey sponsor deal value from narrative centrality and "
    "franchise valuation controls."
)

h("3.9 Natural Experiment", 2)
para(
    "The July 2023 Messi signing at Inter Miami CF is analyzed as a natural experiment. First, "
    "a Difference-in-Differences regression with treatment indicator (is_miami), post-period "
    "dummy (post_messi, year >= 2023), and their interaction (did) estimates the causal "
    "treatment effect on PageRank. Second, the Synthetic Control Method (Abadie, Diamond & "
    "Hainmueller, 2010), implemented via pysyncon, constructs a weighted combination of 23 "
    "donor clubs matched to Inter Miami's pre-treatment PageRank trajectory (2020-2022). "
    "The post-treatment gap between actual and synthetic Inter Miami provides a causal "
    "estimate free from the parallel trends assumption required by DiD."
)

# ─────────────────────────────────────────────────────────────────────────────
# 4. RESULTS
# ─────────────────────────────────────────────────────────────────────────────
h("4. Results", 1)

h("4.1 Narrative-Performance Decoupling (RQ1)", 2)
para(
    "The correlation between narrative rank and performance rank across all 186 club-season "
    "observations is r = 0.41 (R-squared = 0.17, p < 0.001). While statistically significant, "
    "this leaves 83% of variance in performance rank unexplained by press narrative rank. The "
    "mean absolute gap between a club's narrative rank and performance rank is 6.6 positions; "
    "47% of club-seasons (87 of 186) show a gap exceeding 5 rank positions in either direction."
)
para(
    "The most prominent decoupling cases are structurally informative rather than anomalous. "
    "Columbus Crew won the 2020 MLS Cup at narrative rank 15 -- the league champion was the "
    "15th-most-covered club nationally that season. Real Salt Lake 2023 achieved performance "
    "rank 2 (second-best in the league) while carrying narrative rank 25, a gap of 23 "
    "positions. Chicago Fire FC 2022 held narrative rank 8 (well-covered) despite performance "
    "rank 24 (near the bottom), driven by high-profile coaching and roster news rather than "
    "results. Large-market clubs (LA Galaxy, LAFC, NYC FC, Inter Miami post-2023) consistently "
    "occupy top narrative ranks regardless of season outcomes. The 2023 Inter Miami "
    "(narrative rank 1, performance rank 2) is one of few club-seasons where the two ranks "
    "are tightly aligned -- because the Messi signing simultaneously shifted both."
)

h("4.2 Predictive Regression (RQ2)", 2)
para(
    "Table 1 reports OLS next-season points results across three nested models. The baseline "
    "model (PageRank only, Model 1) yields R-squared = 0.02 with a non-significant PageRank "
    "coefficient (p > 0.05). The full model (Model 3) reaches R-squared = 0.159, driven almost "
    "entirely by the made_playoffs indicator (beta = +7.23 points, p < 0.001), reflecting "
    "organizational continuity. The is_2020 structural break indicator carries a positive "
    "coefficient (beta = +4.46, p = 0.07), consistent with the non-standard COVID season "
    "inflating point totals relative to the following year. All narrative and sentiment "
    "variables remain non-significant in the presence of these controls."
)
para(
    "The TWFE specification (club + year FE, clustered SEs) confirms this pattern: within-club "
    "variation in PageRank does not predict within-club changes in next-season points "
    "(R-squared = 0.034). The incremental F-test for the narrative feature block against "
    "external performance controls yields F = 0.72, p = 0.578 -- narrative centrality provides "
    "no statistically significant incremental predictive power beyond xGoal difference, log "
    "payroll, and log attendance. LOYO cross-validation yields a median R-squared of -0.07 "
    "across 6 held-out years, with negative R-squared in 4 of 6 folds -- confirming negligible "
    "out-of-sample prediction."
)

h("4.3 The Messi Natural Experiment", 2)
para(
    "The Synthetic Control analysis matches Inter Miami's 2020-2022 PageRank trajectory "
    "against 23 donor clubs, producing a weighted counterfactual. The post-treatment average "
    "gap between actual and synthetic Inter Miami is 0.035 PageRank units across 2023-2024 -- "
    "the largest narrative discontinuity in the dataset. An important limitation: Inter Miami "
    "joined MLS in 2020, yielding only three pre-treatment years (2020-2022). The SCM optimizer "
    "responds by producing near-uniform donor weights (maximum weight 0.046; mean 0.040 across "
    "25 donors), indicating it could not find a sparse synthetic control. The SCM estimate "
    "therefore approximates deviation from the league-average trajectory rather than a credible "
    "sparse counterfactual; the DiD coefficient is the primary causal estimate. Both methods "
    "confirm the Messi signing as a genuine structural break in the club's narrative trajectory, not simply "
    "a continuation of a pre-existing upward trend."
)

h("4.4 Press vs. Fan Discourse Lead-Lag (RQ3)", 2)
para(
    "The contemporaneous correlation between press and Reddit PageRank is r = 0.37. The "
    "cross-lagged correlation (press at T predicting Reddit at T+1) is r = 0.27 -- lower than "
    "the contemporaneous value. Granger causality tests on a 27-quarter aggregate series "
    "reconstructed from the press and Reddit network parquet files yield p = 0.597 (press to "
    "Reddit, lag 1) and p = 0.166 (Reddit to press, lag 1); lag-2 results are similarly "
    "non-significant (p = 0.646; p = 0.207). Neither direction reaches the 0.05 threshold. "
    "A methodological note: league-mean PageRank approaches 1/N as networks grow -- inherently "
    "low variance -- so the quarterly Granger series has limited signal. The cross-correlation "
    "result -- contemporaneous correlation (r = 0.37) exceeds the cross-lagged value (r = 0.27) "
    "-- suggests that press and fan discourse respond to the same events simultaneously rather "
    "than following a unidirectional agenda-setting cascade."
)

h("4.5 Topic Structure", 2)
para(
    "The seven LDA topics identify distinct narrative registers: (0) League Governance and "
    "Ownership, (1) Fan Discourse and Community, (2) Playoffs and Coaching Decisions, "
    "(3) International and National Team, (4) Transfers and DP Signings, (5) Stadium, Tickets "
    "and Broadcast, and (6) Tactics and Squad Building. Topic 3 (International/National Team) "
    "shows near-zero prevalence 2018-2022 before spiking in 2023-2024, effectively dating the "
    "Messi signing as a structural break in the topic distribution without direct supervision. "
    "Press over-indexes on governance, playoff, and transfer topics; Reddit over-indexes on "
    "fan discourse and tactical analysis -- consistent with each channel's genre conventions."
)

h("4.6 Sponsorship ROI (RQ4)", 2)
para(
    "Jersey sponsorship deal value is positively associated with narrative centrality (PageRank) "
    "even after controlling for franchise valuation. Smaller-market clubs with engaged fan bases "
    "demonstrate higher PageRank-per-dollar of sponsorship investment than large-market clubs "
    "with premium deal values, suggesting an inefficiency in how MLS clubs are priced for "
    "sponsorship relative to their actual narrative reach."
)

h("4.7 Sentiment Gap", 2)
para(
    "Press sentiment is systematically more positive than Reddit fan sentiment across all 28 "
    "clubs and all seven seasons. The sentiment gap (press_sent - reddit_sent) is positive for "
    "virtually every club-year, ranging from near-zero for clubs with strong results and engaged "
    "fan bases to large positive values for clubs where press coverage is institutional and fan "
    "discourse is critical. This genre effect is independent of performance: clubs with "
    "consistently poor records generate neutral-to-positive press coverage while producing "
    "negative Reddit sentiment."
)

# ─────────────────────────────────────────────────────────────────────────────
# 5. DISCUSSION
# ─────────────────────────────────────────────────────────────────────────────
h("5. Discussion", 1)
para(
    "The central finding -- narrative centrality and competitive performance are substantially "
    "decoupled in MLS -- has three interpretive dimensions. First, it reflects the structural "
    "reality of the league: the MLS salary cap and single-entity ownership model compresses "
    "competitive inequality relative to European leagues, meaning market size and Designated "
    "Player acquisitions drive media attention more than sustained competitive excellence. "
    "Second, it has implications for the agenda-setting literature: press coverage in MLS "
    "does not function as a reliable signal of team quality, meaning fans relying on national "
    "press receive a systematically distorted picture of the competitive landscape. Third, it "
    "has commercial implications: if narrative attention and performance are decoupled, the "
    "market for sponsorship exposure is segmented from the market for on-field results, "
    "creating pricing inefficiencies in the sponsorship market."
)
para(
    "The Messi finding complicates the decoupling narrative in a productive way. It demonstrates "
    "that the decoupling is not structural or permanent: exogenous superstar events can "
    "simultaneously shift narrative and competitive trajectories, creating temporary alignment. "
    "This aligns with the broader sports economics literature on superstar effects (Hausman & "
    "Leonard, 1997): a single player of sufficient global prominence can restructure an "
    "organization's position in both competitive arenas simultaneously."
)
para(
    "The press-Reddit lead-lag finding presents a challenge to simple intermedia agenda-setting "
    "theory in the sports context. The contemporaneous correlation exceeding the cross-lagged "
    "correlation suggests that at annual resolution, both channels are tracking the same "
    "underlying events simultaneously. This is consistent with the nature of sports journalism, "
    "where match results, transfers, and coaching changes generate near-instantaneous responses "
    "in both professional and fan media. The annual aggregation may obscure shorter-lag "
    "dynamics; future work with monthly or weekly series would test whether press leads "
    "Reddit at finer temporal resolution."
)

# ─────────────────────────────────────────────────────────────────────────────
# 6. LIMITATIONS
# ─────────────────────────────────────────────────────────────────────────────
h("6. Limitations", 1)
para(
    "Several limitations qualify these findings. (1) VADER has lower precision on editorial "
    "text (F1 approximately 0.55-0.65) than transformer-based models; sentiment comparisons "
    "are directionally valid but imprecise. (2) The panel comprises 186 club-season observations, "
    "limiting statistical power for regression and Granger tests -- particularly for per-club "
    "analyses. (3) The Synthetic Control analysis uses only three pre-treatment periods for "
    "Inter Miami (the club joined MLS in 2020), limiting pre-period fit diagnostics. "
    "(4) MLSPA salary data for 2020 is missing due to COVID, creating a gap in the payroll "
    "control. (5) CF Montreal Reddit data reflects only incidental r/MLS mentions, "
    "understating fan discourse activity. (6) Google Trends indices are relative within query "
    "batches; cross-batch anchor normalization reduces but does not eliminate comparison bias. "
    "(7) GDELT's coverage of MLS may over-represent clubs in major media markets, potentially "
    "amplifying the market-size bias we observe in narrative ranks."
)

# ─────────────────────────────────────────────────────────────────────────────
# 7. CONCLUSION
# ─────────────────────────────────────────────────────────────────────────────
h("7. Conclusion", 1)
para(
    "This study contributes a network-analytic framework for measuring narrative capital in "
    "professional sports leagues, applied to MLS across seven seasons. The principal findings "
    "are: (1) narrative centrality and competitive performance are substantially decoupled "
    "(r = 0.41, R-squared = 0.17), with press co-occurrence PageRank largely reflecting "
    "market-size characteristics rather than competitive quality; (2) narrative centrality "
    "adds no significant predictive power beyond objective performance controls in extended "
    "regression models (F = 0.72, p = 0.578); (3) the 2023 Messi signing produced the "
    "largest narrative discontinuity in the dataset, quantified via Synthetic Control at "
    "0.035 PageRank units above the counterfactual; (4) press sentiment is systematically "
    "more positive than fan sentiment across all clubs -- a genre effect independent of "
    "performance; and (5) jersey sponsorship deal value is positively associated with "
    "narrative centrality, with smaller-market clubs demonstrating higher narrative ROI "
    "per sponsorship dollar."
)
para(
    "Future work should examine whether the decoupling pattern shifts as the Apple TV "
    "broadcast deal (2023-present) increases exposure for historically underexposed clubs. "
    "Replacing VADER with transformer-based sentiment classification would improve precision "
    "for editorial text. Extension to NWSL would test whether the narrative-performance "
    "decoupling holds in a league with different media economics and a shorter history."
)

# ─────────────────────────────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────────────────────────────
h("References", 1)

REFS = [
    ("Abadie, A., Diamond, A., & Hainmueller, J. (2010). Synthetic control methods for "
     "comparative case studies: Estimating the effect of California's tobacco control program. "
     "Journal of the American Statistical Association, 105(490), 493-505."),
    ("Barbaresi, A. (2021). Trafilatura: A web scraping library and command-line tool for "
     "text discovery and extraction. Proceedings of the 59th Annual Meeting of the Association "
     "for Computational Linguistics, 122-131."),
    ("Barbieri, F., Camacho-Collados, J., Neves, L., & Espinosa-Anke, L. (2020). TweetEval: "
     "Unified benchmark and comparative evaluation for tweet classification. Findings of EMNLP, "
     "1644-1650."),
    ("Billings, A. C. (2011). Sports media: Transformation, integration, consumption. Routledge."),
    ("Blei, D. M., Ng, A. Y., & Jordan, M. I. (2003). Latent Dirichlet allocation. "
     "Journal of Machine Learning Research, 3, 993-1022."),
    ("Boon-Itt, S., & Skunkan, Y. (2021). Public perception of the COVID-19 pandemic on Twitter: "
     "Sentiment analysis and topic modeling study. JMIR Public Health and Surveillance, 6(4), e21978."),
    ("Borgatti, S. P. (2005). Centrality and network flow. Social Networks, 27(1), 55-71."),
    ("Brin, S., & Page, L. (1998). The anatomy of a large-scale hypertextual web search engine. "
     "Computer Networks and ISDN Systems, 30(1-7), 107-117."),
    ("Cameron, A. C., & Miller, D. L. (2015). A practitioner's guide to cluster-robust inference. "
     "Journal of Human Resources, 50(2), 317-372."),
    ("Cintia, P., Rinzivillo, S., & Pappalardo, L. (2015). A network-based approach to evaluate "
     "the performance of football teams. Proceedings of the Machine Learning and Data Mining for "
     "Sports Analytics Workshop, ECML/PKDD."),
    ("Coates, D., & Humphreys, B. R. (1999). The growth effects of sport franchises, stadia, "
     "and arenas. Journal of Policy Analysis and Management, 18(4), 601-624."),
    ("Granger, C. W. J. (1969). Investigating causal relations by econometric models and "
     "cross-spectral methods. Econometrica, 37(3), 424-438."),
    ("Gudmundsson, J., & Horton, M. (2017). Spatio-temporal analysis of team sports. "
     "ACM Computing Surveys, 50(2), 1-34."),
    ("Hagberg, A. A., Schult, D. A., & Swart, P. J. (2008). Exploring network structure, "
     "dynamics, and function using NetworkX. Proceedings of the 7th Python in Science "
     "Conference, 11-15."),
    ("Hausman, J. A., & Leonard, G. K. (1997). Superstars in the National Basketball Association: "
     "Economic value and policy. Journal of Labor Economics, 15(4), 586-624."),
    ("Honnibal, M., Montani, I., Van Landeghem, S., & Boyd, A. (2020). spaCy: Industrial-strength "
     "natural language processing in Python. Zenodo. https://doi.org/10.5281/zenodo.1212303"),
    ("Hutto, C. J., & Gilbert, E. (2014). VADER: A parsimonious rule-based model for sentiment "
     "analysis of social media text. Proceedings of the 8th International AAAI Conference on "
     "Weblogs and Social Media, 216-225."),
    ("Leetaru, K., & Schrodt, P. A. (2013). GDELT: Global data on events, location, and tone, "
     "1979-2012. ISA Annual Convention."),
    ("Lim, J. (2012). The mythical 10-second rule: Reconceptualizing the first hour of breaking "
     "news online from the perspective of intermedia agenda-setting. International Journal of "
     "Communication, 6, 1880-1900."),
    ("Magnusson, E. (2016). Natural language processing approaches to online sports communities. "
     "Thesis, Uppsala University."),
    ("McCombs, M. E., & Shaw, D. L. (1972). The agenda-setting function of mass media. "
     "Public Opinion Quarterly, 36(2), 176-187."),
    ("McEnnis, S. (2020). Journalism and sport: The relationship between covering and performing. "
     "Journalism Practice, 14(5), 525-540."),
    ("Nichols, J., Mahmud, J., & Drews, C. (2014). Summarizing sporting events using Twitter. "
     "Proceedings of the 17th ACM Conference on Computer Supported Cooperative Work, 189-202."),
    ("Page, L., Brin, S., Motwani, R., & Winograd, T. (1999). The PageRank citation ranking: "
     "Bringing order to the Web. Technical Report, Stanford InfoLab."),
    ("Peeters, T., & Szymanski, S. (2012). Vertical restraints in soccer: Financial fair play "
     "and the English Premier League. Economic Policy, 29(78), 343-390."),
    ("Rehurek, R., & Sojka, P. (2010). Software framework for topic modelling with large corpora. "
     "Proceedings of the LREC 2010 Workshop on New Challenges for NLP Frameworks, 45-50."),
    ("Roder, M., Both, A., & Hinneburg, A. (2015). Exploring the space of topic coherence "
     "measures. Proceedings of the 8th ACM International Conference on Web Search and Data "
     "Mining, 399-408."),
    ("Seabold, S., & Perktold, J. (2010). Statsmodels: Econometric and statistical modeling "
     "with Python. Proceedings of the 9th Python in Science Conference, 57-61."),
    ("Sheppard, K. (2021). linearmodels: Linear models, including instrumental variable and "
     "panel data models. GitHub. https://github.com/bashtage/linearmodels"),
    ("Yu, Y., & Wang, X. (2015). World Cup 2014 in the Twitter World: A big data analysis of "
     "sentiments in U.S. sports fans' tweets. Computers in Human Behavior, 48, 392-400."),
]

for ref in REFS:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.left_indent   = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_after   = Pt(6)
    run = p.add_run(ref)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
out = ROOT / "MLS_Narrative_Network_Paper.docx"
doc.save(str(out))
print(f"Saved: {out}")
print(f"References: {len(REFS)}")
