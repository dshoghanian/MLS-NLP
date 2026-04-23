"""
Generate Professor_Meeting_Doc.docx — casual talking points for professor meeting.
Run: /opt/anaconda3/bin/python scripts/build_meeting_doc.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

def h(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        if color:
            run.font.color.rgb = color
    return p

def body(text, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    return p

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("→  " + text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def bullet(text, bold_start=None, level=0, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.2*level)
    p.paragraph_format.space_after = Pt(3)
    if bold_start:
        r1 = p.add_run(bold_start + " ")
        r1.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.font.name = "Calibri"; r2.font.size = Pt(size)
    return p

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

# ─────────────────────────────────────────────────────────────────────────────
# TITLE
# ─────────────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Professor Meeting — What to Say\n")
r.bold = True; r.font.size = Pt(18); r.font.name = "Calibri"
r2 = title.add_run("MLS Narrative Network Project  |  April 2026")
r2.font.size = Pt(12); r2.font.name = "Calibri"; r2.italic = True

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# THE GIST — open with this
# ─────────────────────────────────────────────────────────────────────────────
h("The Gist — Open With This", level=1)
body(
    "If he asks 'so what are you doing?' — say this:"
)
tip(
    "We're using network analysis and NLP on 7 years of MLS press coverage and Reddit fan posts "
    "to measure how much of the media conversation each club owns. We call it narrative centrality. "
    "The main question is: does being talked about more actually help you win? Short answer — not really. "
    "But that null result is interesting because it means media attention and competitive performance "
    "are separate things, and clubs can strategically manage one without the other."
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
# WHAT WE ACTUALLY BUILT
# ─────────────────────────────────────────────────────────────────────────────
h("What We Actually Built", level=1)
bullet("Collected ~7,200 press articles (GDELT + club RSS feeds) and ~16,000 Reddit posts across 28 MLS clubs, 2018–2024.", level=0)
bullet("Built co-occurrence networks — if two clubs get mentioned in the same article, they're connected. Ran Google PageRank on the network to rank each club's narrative importance.", level=0)
bullet("Ran VADER sentiment on all articles and Reddit posts to get press vs. fan tone for each club each year.", level=0)
bullet("Pulled in external data: ASA xGoals, MLSPA salary records, Forbes valuations, home attendance, Google Trends — all merged into one big panel dataset.", level=0)
bullet("Built a Jupyter notebook (00_showcase.ipynb) that runs the entire analysis end to end — all the regressions, visualizations, everything.", level=0)
bullet("Also built a Streamlit app (app.py) for interactive exploration of the network.", level=0)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
# HIS OUTLINE — answer each section
# ─────────────────────────────────────────────────────────────────────────────
h("Answers to His PDF Outline", level=1)

# INTRO
h("Introduction", level=2)
body("He wants you to be able to answer: what do you do, why does it matter, what's new.")

bullet("What we do:", bold_start="What we do:", level=0)
tip("We build network graphs from news articles and Reddit posts, compute PageRank for each MLS club per season, and test whether that score predicts future performance, sentiment, or sponsorship value.")

bullet("Why it matters:", bold_start="Why it matters:", level=0)
tip("Clubs and sponsors spend a lot of money on media strategy assuming it drives value. We give them a data-driven way to actually measure their narrative footprint — not just count mentions but understand position in the broader conversation.")

bullet("What's new (say at least two of these):", bold_start="What's new:", level=0)
bullet("First study to compare press AND Reddit narratives for the same league at this scale using network analysis.", level=1)
bullet("PageRank applied to sports media co-occurrence — hadn't been done for MLS.", level=1)
bullet("We formally test whether narrative attention predicts winning — turns out it mostly doesn't, which is itself a finding.", level=1)
bullet("We treat the Messi signing as a natural experiment with a real causal method (Difference-in-Differences + Synthetic Control).", level=1)
spacer()

# FINDINGS SUMMARY
h("Quick Summary of What We Found", level=2)
body("Know these numbers cold — he will probably ask.")

bullet("Narrative rank and performance rank are correlated at r = 0.41. Moderate — clubs that get more press tend to do better, but it's far from a clean relationship.", level=0)
bullet("The average club is off by 6.6 rank positions between how much press they get vs. how well they actually perform.", level=0)
bullet("In every regression model we ran, narrative variables did NOT significantly predict next-season points. The only thing that predicted future performance was whether you made the playoffs last year (β = +7 points, p < 0.001). That's the playoff persistence effect — good teams tend to stay good.", level=0)
bullet("When we added xGoals, payroll, and attendance as controls, narrative features added zero extra predictive power (F = 0.72, p = 0.578).", level=0)
bullet("LOYO cross-validation (leave one year out) gave median R² = −0.07. Negative. The model predicts worse than just guessing the average. So it really doesn't generalize out of sample.", level=0)
bullet("Messi: Inter Miami's PageRank jumped +0.035 units above the synthetic counterfactual after the signing. Biggest narrative discontinuity in the entire dataset.", level=0)
bullet("Press is consistently more positive than Reddit fans (mean gap ~+0.18). Chicago Fire FC has the worst gap — journalists are neutral, fans are very negative.", level=0)
bullet("Sponsorship deal value goes up with narrative centrality. Smaller market clubs that punch above their weight in media get better ROI per sponsorship dollar.", level=0)
spacer()

# THEORY
h("Theory Section — What to Say", level=2)
body("He listed three areas. Have a sentence or two for each.")

bullet("Agenda-setting theory (McCombs & Shaw 1972): the idea that what the media covers shapes what the public thinks is important. We're testing whether that flows from press to Reddit fans. Turns out both channels react to the same events at the same time — no clear leader.", bold_start="Predictive power of press and social media:", level=0)
bullet("Sports analytics is all xGoals and advanced stats. We're adding the media layer on top and asking if it adds anything predictive. Answer: not really for performance, but yes for sponsorship value.", bold_start="Predictive practices in sports:", level=0)
bullet("PageRank from Google — same algorithm that ranks web pages — applied to who-mentions-who in sports articles. Gives a better picture of narrative importance than just raw mention counts because it accounts for who you're mentioned alongside.", bold_start="Network analysis:", level=0)
spacer()

# METHOD
h("Methods — What to Be Ready to Explain", level=2)

bullet("Co-occurrence network: count how many articles mention Club A and Club B together. That count is the edge weight. Build that graph for every year (and every quarter for time-series tests).", bold_start="Network generation:", level=0)
bullet("Three nested OLS models predicting next-season points. Also a Two-Way Fixed Effects (TWFE) model with club and year fixed effects — that's the gold standard for panel data because it absorbs all the stuff that doesn't change per club (market size, history, etc.) and all year-level shocks (COVID, expansion, etc.).", bold_start="Linear regression:", level=0)
bullet("We also ran logistic regression predicting playoff qualification, Leave-One-Year-Out cross-validation, and an external controls model with xGoals + payroll + attendance.", bold_start="Other models:", level=0)
bullet("For Messi: DiD regression (is_miami × post_2023 interaction) and Synthetic Control (weighted average of 23 donor clubs as a counterfactual).", bold_start="Causal methods:", level=0)
bullet("LDA topic modeling found 7 themes in the combined corpus. We had to add domain-specific stopwords (removed NBA/NFL/NHL terms that were contaminating the topics).", bold_start="Topic modeling:", level=0)
bullet("Granger causality test on quarterly press vs. Reddit PageRank series (N=27 quarters). Upgraded from annual (N=7, too small to trust) to quarterly. Still not significant in either direction.", bold_start="Granger:", level=0)
spacer()

# DISCUSSION
h("Discussion — Theoretical and Practical Contributions", level=2)
body("He specifically asked about these two. Have an answer for each.")

bullet(
    "We're showing that narrative capital is a real, measurable construct that lives partly "
    "independently of winning. That's a contribution to sports media theory — it means clubs can "
    "build (or lose) narrative standing for reasons unrelated to results.",
    bold_start="Theoretical — value of media:", level=0
)
bullet(
    "We're offering a framework that any league or club can run on public data. "
    "Press + Reddit + PageRank + sentiment gap = a monitoring dashboard for your media position. "
    "That's new.",
    bold_start="Theoretical — new framework:", level=0
)
bullet(
    "Clubs can track where they sit on the narrative-performance spectrum. If you're overexposed "
    "(lots of press, bad results) you have a credibility problem. If you're underexposed (winning "
    "but no media) you're probably leaving sponsorship money on the table.",
    bold_start="Practical — managing centrality:", level=0
)
bullet(
    "Honest answer: narrative doesn't predict next-season performance on its own. But "
    "within-season momentum shifts (monthly PageRank trends) can signal organizational changes "
    "before they show in the standings. And combined with xGoals data, it gives a fuller picture "
    "of where a club is headed.",
    bold_start="Practical — predict next season earlier:", level=0
)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
# THINGS TO BE HONEST ABOUT
# ─────────────────────────────────────────────────────────────────────────────
h("Things to Be Ready to Defend / Be Honest About", level=1)

bullet("VADER sentiment isn't the most accurate tool for editorial sports journalism. It works better on social media text. We use it as a directional signal.", level=0)
bullet("The Synthetic Control for Messi is limited — Inter Miami only joined MLS in 2020 so we had 3 pre-treatment years. The optimizer spread weights almost equally across all donor clubs, meaning it's basically using the league average as a counterfactual, not a sparse matched control. DiD is our real causal estimate.", level=0)
bullet("The panel is small (186 club-seasons). TWFE absorbs 35 fixed effects (28 clubs + 7 years) which doesn't leave much room for the narrative variables to show significance.", level=0)
bullet("The main finding — narrative doesn't predict performance — sounds like a negative result. But frame it as: the decoupling of narrative capital from competitive performance is the finding. It's not a failed attempt to predict wins, it's a demonstration that these are different dimensions.", level=0)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
# QUESTIONS TO ASK HIM
# ─────────────────────────────────────────────────────────────────────────────
h("Questions to Ask Him", level=1)

bullet("The null predictive result is our strongest, most rigorous finding. Is that enough of a contribution for the target journals or do we need to reframe the paper around a different angle?", level=0)
bullet("He mentioned Managing Sport and Leisure as a target — it publishes shorter research notes with faster review. Should we aim for that format instead of a full paper?", level=0)
bullet("The theory section needs more depth — which literature does he think we're most clearly speaking to? Agenda-setting, sports media, or sports analytics?", level=0)
bullet("Is there a co-author from the sport management side he'd recommend to strengthen the framing?", level=0)
bullet("Does he want us to extend the dataset back before 2018 or is 2018–2024 sufficient?", level=0)
spacer()

# ─────────────────────────────────────────────────────────────────────────────
# KEY NUMBERS CHEAT SHEET
# ─────────────────────────────────────────────────────────────────────────────
h("Key Numbers — Cheat Sheet", level=1)
body("Stick this at the end. If you go blank mid-meeting, flip here.")

rows = [
    ("Press articles collected", "~7,200"),
    ("Reddit posts collected", "~15,800"),
    ("Clubs in panel", "28"),
    ("Years covered", "2018–2024 (7 seasons)"),
    ("Club-season obs for regression", "186"),
    ("Narrative rank ~ Performance rank", "r = 0.41"),
    ("Average narrative-performance gap", "6.6 rank positions"),
    ("Full OLS model R²", "0.159"),
    ("Only significant predictor", "made_playoffs  β = +7.2 pts  p < 0.001"),
    ("LOYO median R²", "−0.07  (model doesn't generalize)"),
    ("Incremental F-test (narrative vs. perf controls)", "F = 0.72,  p = 0.578  (no added power)"),
    ("Messi SCM post-treatment gap", "+0.035 PageRank units above counterfactual"),
    ("Granger test series length", "27 quarters  (upgraded from 7 annual)"),
    ("Granger press → Reddit (lag 1)", "p = 0.597  (not significant)"),
    ("Mean sentiment gap (press − Reddit)", "~+0.18  (press always more positive)"),
    ("LDA topics", "7 topics, retrained with domain stopwords"),
]

table = doc.add_table(rows=1 + len(rows), cols=2)
table.style = 'Light Shading Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Metric"; hdr[1].text = "Value"
for cell in hdr:
    for par in cell.paragraphs:
        for run in par.runs:
            run.bold = True; run.font.name = "Calibri"; run.font.size = Pt(10)

for row_data, row in zip(rows, table.rows[1:]):
    row.cells[0].text = row_data[0]
    row.cells[1].text = row_data[1]
    for cell in row.cells:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.name = "Calibri"; run.font.size = Pt(10)

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
out = ROOT / "Professor_Meeting_Doc.docx"
doc.save(out)
print(f"Saved: {out}")
